import os
import sys
import csv
import json
from docx import Document
from docx.shared import Inches

# Install the necessary library: pip install python-docx

# This script generates a Low-Level Design (LLD) document in .docx format
# using the outputs from the discovery, mapping, and IaC generation steps.

def main(gcp_resources_file, aws_mapping_file, aws_tf_file, output_docx_file):
    """
    Generates the LLD document.
    """
    if not all(os.path.exists(f) for f in [gcp_resources_file, aws_mapping_file, aws_tf_file]):
        print("One or more input files not found.", file=sys.stderr)
        sys.exit(1)

    print("[*] Starting LLD document generation...")
    document = Document()
    document.add_heading('AWS Migration - Low-Level Design Document', 0)
    document.add_paragraph('This document outlines the design for the migrated infrastructure in AWS.')

    # --- Section 1: Resource Inventory ---
    document.add_heading('1. GCP Resource Inventory', level=1)
    with open(gcp_resources_file, 'r') as f:
        gcp_resources = json.load(f)
    
    document.add_paragraph(f"Total GCP resources discovered: {len(gcp_resources)}")
    for resource in gcp_resources[:20]: # Show a sample
        document.add_paragraph(f"  - {resource}", style='List Bullet')
    if len(gcp_resources) > 20:
        document.add_paragraph("  - ...and more.", style='List Bullet')

    # --- Section 2: Service Mapping ---
    document.add_heading('2. Service Mapping (GCP to AWS)', level=1)
    table = document.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'GCP Resource'
    hdr_cells[1].text = 'AWS Equivalent'
    hdr_cells[2].text = 'Details'

    with open(aws_mapping_file, 'r') as f:
        reader = csv.reader(f)
        next(reader) # Skip header
        for row in reader:
            if len(row) == 3:
                cells = table.add_row().cells
                cells[0].text = row[0]
                cells[1].text = row[1]
                cells[2].text = row[2]

    # --- Section 3: AWS Terraform Configuration ---
    document.add_heading('3. AWS Infrastructure as Code', level=1)
    document.add_paragraph('The following Terraform code will be used to provision the AWS infrastructure.')
    
    with open(aws_tf_file, 'r') as f:
        terraform_code = f.read()

    document.add_paragraph('```hcl\n' + terraform_code + '\n```', style='No Spacing')

    print("[*] Saving document...")
    document.save(output_docx_file)
    print(f"[✓] LLD document generated at: {output_docx_file}")


if __name__ == "__main__":
    if len(sys.argv) < 5:
        print("Usage: python3 generate_aws_lld.py <GCP_RESOURCES_JSON> <AWS_MAPPING_CSV> <AWS_TF_FILE> <OUTPUT_DOCX_FILE>")
        sys.exit(1)

    gcp_res_file = sys.argv[1]
    aws_map_file = sys.argv[2]
    aws_tf_file = sys.argv[3]
    output_file = sys.argv[4]
    main(gcp_res_file, aws_map_file, aws_tf_file, output_file)
