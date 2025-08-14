from docx import Document
import os

def fill_lld_template(template_path, out_path, context):
    doc = Document(template_path)
    # naive replace: go through paragraphs and replace placeholders like {{VPC_CIDR}}
    for p in doc.paragraphs:
        for k,v in context.items():
            if f"{{{{{k}}}}}" in p.text:
                p.text = p.text.replace(f"{{{{{k}}}}}", str(v))
    # tables too
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for k,v in context.items():
                    if f"{{{{{k}}}}}" in cell.text:
                        cell.text = cell.text.replace(f"{{{{{k}}}}}", str(v))
    doc.save(out_path)
